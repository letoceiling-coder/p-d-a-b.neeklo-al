# -*- coding: utf-8 -*-
from docx import Document

d = Document()
d.add_paragraph("ИНН: 7701234567")
d.add_paragraph("Сумма контракта: 120500 рублей")
d.add_paragraph("Дата начала: 12.03.2026")
d.add_paragraph("Дата окончания: 12.04.2026")
d.save("/tmp/test-rule.docx")
print("ok")
